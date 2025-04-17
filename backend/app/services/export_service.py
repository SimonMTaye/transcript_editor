from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Dict
import os


class ExportService:
    def __init__(self):
        self.export_dir = "exports"
        os.makedirs(self.export_dir, exist_ok=True)

    def export_to_word(self, transcript: Dict, timestamps: List[Dict] = None) -> str:
        """Export transcript to Word document with optional timestamps"""
        doc = Document()

        # Add title
        title = doc.add_heading(transcript["title"], level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add content
        if timestamps:
            for segment in timestamps:
                p = doc.add_paragraph()
                if transcript["status"] != "raw":
                    # Add timestamp for non-raw transcripts
                    p.add_run(
                        f"[{segment['start']:.2f}s - {segment['end']:.2f}s] "
                    ).italic = True
                p.add_run(segment["text"])
        else:
            # For transcripts without timestamps
            paragraphs = transcript["content"].split("\n")
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())

        # Save the document
        filename = f"transcript_{transcript['id']}_{transcript['status']}.docx"
        filepath = os.path.join(self.export_dir, filename)
        doc.save(filepath)

        return filepath


export_service = ExportService()
